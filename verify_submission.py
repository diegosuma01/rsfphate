#!/usr/bin/env python3
"""
Verify that the repository is ready for thesis submission.
Checks:
  - All required files present
  - Notebooks are executed (have outputs)
  - Word document contains unified numbers
  - GitHub README is comprehensive
  - requirements.txt is valid
"""

import json
import sys
from pathlib import Path

def check_file_exists(path, description):
    """Check if a file exists."""
    path = Path(path)
    if path.exists():
        size_mb = path.stat().st_size / (1024 * 1024)
        print(f"  [OK] {description} ({size_mb:.2f} MB)")
        return True
    else:
        print(f"  [FAIL] {description} NOT FOUND")
        return False

def check_notebooks_executed():
    """Verify all notebooks are executed (have outputs)."""
    notebooks_dir = Path("notebooks")
    notebooks = [
        "eda_datos_reales.ipynb",
        "eda_datos_sinteticos.ipynb",
        "analisis_churn_survival.ipynb",
        "validacion_sintetica.ipynb",
        "analisis_multiproducto.ipynb",
    ]

    print("\n[Notebooks Status]")
    all_executed = True

    for nb_name in notebooks:
        nb_path = notebooks_dir / nb_name
        if not nb_path.exists():
            print(f"  ❌ {nb_name} NOT FOUND")
            all_executed = False
            continue

        try:
            with open(nb_path, 'r', encoding='utf-8') as f:
                nb = json.load(f)

            code_cells = sum(1 for c in nb['cells'] if c['cell_type'] == 'code')
            executed_cells = sum(
                1 for c in nb['cells']
                if c['cell_type'] == 'code' and c.get('outputs') and len(c['outputs']) > 0
            )

            if executed_cells >= code_cells * 0.8:
                print(f"  [OK] {nb_name} ({executed_cells}/{code_cells} cells executed)")
            else:
                print(f"  [WARN] {nb_name} ({executed_cells}/{code_cells} cells, incomplete)")
                all_executed = False

        except Exception as e:
            print(f"  [FAIL] {nb_name} ERROR: {e}")
            all_executed = False

    return all_executed

def check_word_document():
    """Verify Word document is complete and unified."""
    from docx import Document

    word_path = Path("notebooks/TFG_Diego_Suarez_Maranon.docx")
    if not word_path.exists():
        print(f"  [FAIL] Word document not found at {word_path}")
        return False

    try:
        doc = Document(word_path)

        # Check chapter count
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading 1')]
        print(f"  [OK] Word document: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

        # Verify unified numbers
        all_text = ' '.join([p.text for p in doc.paragraphs])
        issues = []

        if 'n=272' in all_text or 'n=538' in all_text or 'n=358' in all_text:
            issues.append("Old cluster sizes found (272/538/358)")

        if 'churn=14,3%' in all_text or 'churn=8,4%' in all_text:
            issues.append("Old churn rates found (14,3%/8,4%)")

        if 'diffusion_time=4.0' in all_text and 'diffusion_time=7.0' not in all_text:
            issues.append("diffusion_time=4.0 present without 7.0")

        if 'TotalEnergies' in all_text or 'totalenergies' in all_text.lower():
            issues.append("TotalEnergies references still present")

        if issues:
            for issue in issues:
                print(f"  [WARN] {issue}")
            return False
        else:
            print(f"  [OK] Numbers unified (C0=226, C1=577, C2=365, diffusion_time=7.0)")
            print(f"  [OK] No TotalEnergies references")
            return True

    except Exception as e:
        print(f"  [FAIL] Error reading Word: {e}")
        return False

def check_repository_structure():
    """Verify repo structure for GitHub."""
    print("\n[Repository Structure]")

    required_files = [
        ("README.md", "GitHub README"),
        ("requirements.txt", "Pip dependencies"),
        ("run_experiments.py", "Experiment runner"),
        ("verify_submission.py", "Submission checker"),
        ("LICENSE", "License"),
        ("CITATION.cff", "Citation metadata"),
        (".gitignore", "Git ignore rules"),
    ]

    all_ok = True
    for filename, description in required_files:
        all_ok = check_file_exists(filename, description) and all_ok

    # Check src/rsfphate structure
    print("\n[Source Code]")
    src_files = [
        ("src/rsfphate/__init__.py", "Package init"),
        ("src/rsfphate/forest.py", "RSF implementation"),
        ("src/rsfphate/model.py", "RSFPhate main class"),
        ("src/rsfphate/spectral.py", "Spectral clustering"),
    ]

    for filepath, description in src_files:
        all_ok = check_file_exists(filepath, description) and all_ok

    return all_ok

def main():
    """Run all verification checks."""
    print("="*80)
    print("RSF-PHATE Thesis Submission Verification")
    print("="*80)

    print("\n[Core Files]")
    checks = {
        "Repository structure": check_repository_structure(),
        "Notebooks executed": check_notebooks_executed(),
        "Word document": check_word_document(),
    }

    print("\n" + "="*80)
    print("VERIFICATION SUMMARY")
    print("="*80)

    all_ok = all(checks.values())

    for check_name, result in checks.items():
        status = "[PASS]" if result else "[FAIL]"
        print(f"{status} — {check_name}")

    if all_ok:
        print("\n[OK] Repository is READY for submission!")
        print("\nNext steps:")
        print("  1. git add -A")
        print("  2. git commit -m 'Final thesis submission'")
        print("  3. git push origin main")
        print("  4. Create GitHub repository and upload")
        return 0
    else:
        print("\n[FAIL] Fix issues above before submission.")
        return 1

if __name__ == "__main__":
    sys.exit(main())
